import hashlib
import io
import json
import sqlite3
import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from google import genai
import streamlit as st

# ==========================================
# 1. KONFIGURASI HALAMAN & DATABASE SQLITE
# ==========================================
st.set_page_config(
    page_title="LANTIP AI | Lembar Administrasi Pembelajaran Inklusif",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="expanded",
)

DB_NAME = "slb_adminflow.db"


def init_db():
  conn = sqlite3.connect(DB_NAME)
  c = conn.cursor()
  c.execute("""
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT NOT NULL,
            full_name TEXT,
            api_key TEXT,
            nama_dinas TEXT,
            nama_sekolah TEXT,
            alamat_sekolah TEXT,
            nama_kepsek TEXT,
            nip_kepsek TEXT,
            nama_guru TEXT,
            nip_guru TEXT,
            kota_tgl TEXT
        )
    """)
  c.execute("""
        CREATE TABLE IF NOT EXISTS history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            title TEXT,
            mapel TEXT,
            fase_kelas TEXT,
            kekhususan TEXT,
            cp_input TEXT,
            data_json TEXT,
            FOREIGN KEY (username) REFERENCES users (username)
        )
    """)
  conn.commit()
  conn.close()


init_db()


def hash_password(password):
  return hashlib.sha256(str.encode(password)).hexdigest()


def get_user_profile(username):
  conn = sqlite3.connect(DB_NAME)
  c = conn.cursor()
  c.execute("SELECT * FROM users WHERE username = ?", (username,))
  row = c.fetchone()
  conn.close()
  return row


def update_user_profile(username, profile_data):
  conn = sqlite3.connect(DB_NAME)
  c = conn.cursor()
  c.execute(
      """
        UPDATE users SET 
            api_key=?, nama_dinas=?, nama_sekolah=?, alamat_sekolah=?, 
            nama_kepsek=?, nip_kepsek=?, nama_guru=?, nip_guru=?, kota_tgl=?
        WHERE username=?
    """,
      (
          profile_data["api_key"],
          profile_data["nama_dinas"],
          profile_data["nama_sekolah"],
          profile_data["alamat_sekolah"],
          profile_data["nama_kepsek"],
          profile_data["nip_kepsek"],
          profile_data["nama_guru"],
          profile_data["nip_guru"],
          profile_data["kota_tgl"],
          username,
      ),
  )
  conn.commit()
  conn.close()


def save_history(
    username, title, mapel, fase_kelas, kekhususan, cp_input, data_dict
):
  conn = sqlite3.connect(DB_NAME)
  c = conn.cursor()
  data_json = json.dumps(data_dict)
  c.execute(
      """
        INSERT INTO history (username, title, mapel, fase_kelas, kekhususan, cp_input, data_json)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """,
      (
          username,
          title,
          mapel,
          fase_kelas,
          kekhususan,
          cp_input,
          data_json,
      ),
  )
  conn.commit()
  conn.close()


def get_user_history(username):
  conn = sqlite3.connect(DB_NAME)
  c = conn.cursor()
  c.execute(
      """
        SELECT id, created_at, title, mapel, fase_kelas, kekhususan, cp_input, data_json 
        FROM history WHERE username = ? ORDER BY id DESC
    """,
      (username,),
  )
  rows = c.fetchall()
  conn.close()
  return rows


def delete_history_item(item_id):
  conn = sqlite3.connect(DB_NAME)
  c = conn.cursor()
  c.execute("DELETE FROM history WHERE id = ?", (item_id,))
  conn.commit()
  conn.close()


# ==========================================
# 2. STYLING CUSTOM CSS UI MODERN
# ==========================================
st.markdown(
    """
    <style>
    .main-header {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        padding: 24px 30px;
        border-radius: 12px;
        color: white;
        margin-bottom: 20px;
        box-shadow: 0 4px 14px rgba(0,0,0,0.12);
    }
    .main-header h1 { color: #ffffff; font-weight: 800; margin: 0; font-size: 28px; letter-spacing: 0.5px; }
    .main-header p { color: #e0e6ed; margin-top: 6px; margin-bottom: 0; font-size: 13.5px; }
    .info-card {
        background-color: #f8fafc;
        border-left: 4px solid #2563eb;
        padding: 12px 16px;
        border-radius: 6px;
        margin-bottom: 14px;
        font-size: 14px;
    }
    .stButton>button { border-radius: 8px; font-weight: 600; }
    </style>
""",
    unsafe_allow_html=True,
)

# ==========================================
# 3. KONTROL SESSION & APLIKASI
# ==========================================
if "logged_in" not in st.session_state:
  st.session_state.logged_in = False
if "username" not in st.session_state:
  st.session_state.username = ""

if "form_data" not in st.session_state:
  st.session_state.form_data = {
      "elemen": "Bilangan / Analisis Data",
      "ruang_lingkup": (
          "• Pengurutan & Perbandingan Benda\n• Konsep Banyak-Sedikit &"
          " Besar-Kecil"
      ),
      "materi_pokok": (
          "1. Membandingkan banyak-sedikit (1-10)\n2. Mengurutkan benda\n3."
          " Konsep ukuran besar-kecil"
      ),
      "tp_text": (
          "1. Peserta didik mampu membandingkan banyak-sedikit benda konkret"
          " sampai 10.\n2. Peserta didik mampu mengurutkan benda konkret"
          " berdasarkan jumlah.\n3. Peserta didik mampu membedakan ukuran"
          " besar dan kecil suatu benda."
      ),
      "atp1": (
          "1.1 Murid dapat mengamati dua kelompok benda konkret dan membedakan"
          " banyak/sedikit.\n1.2 Murid dapat menunjuk kelompok benda yang lebih"
          " banyak atau sedikit."
      ),
      "atp2": (
          "2.1 Murid dapat membilang benda konkret 1-10 secara berurutan.\n2.2"
          " Murid dapat menyusun benda konkret dari jumlah terkecil ke terbesar."
      ),
      "atp3": (
          "3.1 Murid dapat mengidentifikasi benda berukuran besar dan kecil di"
          " kelas.\n3.2 Murid dapat mengelompokkan benda berdasarkan ukuran."
      ),
      "kegiatan_konkret": (
          "• Eksplorasi benda konkret (balok, kartu PECS, kelereng).\n•"
          " Permainan 'Tebak Ukuran & Jumlah'.\n• Praktik mengelompokkan benda"
          " mandiri."
      ),
      "bentuk_asesmen": (
          "• Unjuk Kerja & Observasi Langsung\n• Lembar Ceklis Perilaku"
          " Pembelajaran\n• Tes Lisan Adaptif"
      ),
      "alokasi_waktu": "12 JP (4 x Pertemuan)",
      "prota_sem1": (
          "1. TP 1: Membandingkan banyak-sedikit benda konkret (1-10) [6 JP]\n2."
          " TP 2: Mengurutkan benda konkret berdasarkan jumlah [6 JP]"
      ),
      "prota_sem2": (
          "1. TP 3: Membedakan ukuran besar dan kecil suatu benda [6 JP]\n2."
          " Evaluasi & Asesmen Sumatif Akhir Tahun [6 JP]"
      ),
      "prosem_detail": (
          "Juli M3-M4: TP 1 (6 JP)\nAgustus M1-M2: TP 2 (6 JP)\nJanuari M2-M3:"
          " TP 3 (6 JP)\nMei M3: Asesmen Sumatif (6 JP)"
      ),
      "dimensi_lulusan": (
          "• Kemandirian dan Ketangguhan\n• Penalaran Kritis dan Pemecahan"
          " Masalah\n• Komunikasi dan Kolaborasi"
      ),
      "modul_ajar": (
          "A. IDENTITAS MODUL\n- Mapel: Matematika Inklusif\n- Fase/Kelas: Fase"
          " A / Kelas I\n- Target Siswa: Hambatan Intelektual\n\nB. LINGKUP 8"
          " DIMENSI PROFIL LULUSAN\n- Kemandirian & Ketangguhan: Berlatih"
          " menunjuk benda secara mandiri.\n- Penalaran Kritis: Membedakan dua"
          " kelompok benda.\n\nC. MEDIA & ALAT BANTU KONKRET\n- Balok kayu"
          " warna-warni, kartu gambar PECS, benda nyata di kelas.\n\nD."
          " LANGKAH PEMBELAJARAN (DEFERENSIASI)\n1. Awal (10 menit): Doa bersama"
          " & Apersepsi bernyanyi angka.\n2. Inti (50 menit): Guru"
          " mendemonstrasikan perbandingan benda. Murid mencoba langsung secara"
          " bergantian.\n3. Penutup (10 menit): Refleksi positif dan pujian"
          " atas kemajuan anak."
      ),
      "ppi_text": (
          "A. PROFIL KEMAMPUAN AWAL (BASELINE)\n- Siswa telah mampu mengenal"
          " simbol angka 1-5 dengan bantuan media visual, namun masih butuh"
          " bimbingan saat membandingkan jumlah benda di atas 5.\n\nB. AKOMODASI"
          " PEMBELAJARAN\n- Menggunakan media konkret taktil dan instrumen"
          " visual PECS.\n- Pendampingan individu secara bertahap (Prompting"
          " Verbal & Fisik).\n\nC. TARGET PROGRAM PEMBELAJARAN INDIVIDUAL"
          " (PPI)\n1. Target Jangka Panjang (1 Tahun): Siswa mampu"
          " membandingkan dan mengurutkan benda konkret 1-10 secara mandiri.\n2."
          " Target Jangka Pendek (Semester 1): Siswa dapat menunjuk kelompok"
          " benda yang lebih banyak atau sedikit pada rentang 1-5."
      ),
      "rubrik_text": (
          "A. CEKLIS OBSERVASI ADAPTIF\n1. Mengamati dua kelompok benda konkret"
          " [Sangat Mampu / Mampu Bantuan / Belum Mampu]\n2. Menunjuk kelompok"
          " benda lebih banyak [Sangat Mampu / Mampu Bantuan / Belum Mampu]\n3."
          " Menyebutkan nama ukuran benda (besar/kecil) [Sangat Mampu / Mampu"
          " Bantuan / Belum Mampu]\n\nB. RUBRIK SKALA PENILAIAN\n- Sangat Mampu"
          " (Skor 3): Melakukan tugas secara mandiri tanpa bantuan.\n- Mampu"
          " dengan Bantuan (Skor 2): Melakukan tugas dengan isyarat/bimbingan"
          " verbal/fisik ringan.\n- Belum Mampu (Skor 1): Membutuhkan bimbingan"
          " penuh dari guru.\n\nC. FORMAT CATATAN ANEKDOT\nTanggal: ........."
          " | Catatan Perilaku / Respon Siswa: ...................."
      ),
  }


# ==========================================
# 4. MODUL HALAMAN LOGIN / REGISTER
# ==========================================
def login_page():
  st.markdown(
      """
        <div class="main-header">
            <h1>🏫 LANTIP AI</h1>
            <p><b>L</b>embar <b>A</b>dmi<b>N</b>is<b>T</b>ras<b>I</b> <b>P</b>embelajaran — Platform Administrasi Guru SLB & Sekolah Inklusif Terintegrasi AI</p>
        </div>
    """,
      unsafe_allow_html=True,
  )

  col1, col2 = st.columns([1, 1])

  with col1:
    st.subheader("🔑 Masuk ke Akun Anda")
    login_user = st.text_input("Username", key="l_user")
    login_pass = st.text_input("Password", type="password", key="l_pass")

    if st.button("🚀 Masuk (Login)", type="primary"):
      conn = sqlite3.connect(DB_NAME)
      c = conn.cursor()
      c.execute(
          "SELECT * FROM users WHERE username = ? AND password = ?",
          (login_user.strip(), hash_password(login_pass)),
      )
      user = c.fetchone()
      conn.close()

      if user:
        st.session_state.logged_in = True
        st.session_state.username = login_user.strip()
        st.success(f"Selamat datang kembali, {user[2] or login_user}!")
        st.rerun()
      else:
        st.error("Username atau Password salah!")

  with col2:
    st.subheader("📝 Buat Akun Guru Baru")
    reg_user = st.text_input("Pilih Username Baru", key="r_user")
    reg_name = st.text_input("Nama Lengkap Guru & Gelar", key="r_name")
    reg_pass = st.text_input("Password Baru", type="password", key="r_pass")

    if st.button("✨ Daftar Akun Baru"):
      if not reg_user or not reg_pass:
        st.warning("Username dan Password wajib diisi!")
      else:
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        try:
          c.execute(
              """
                        INSERT INTO users (username, password, full_name, api_key, nama_dinas, nama_sekolah, alamat_sekolah, nama_kepsek, nip_kepsek, nama_guru, nip_guru, kota_tgl)
                        VALUES (?, ?, ?, '', 'PEMERINTAH DAERAH DIY - DINAS PENDIDIKAN, PEMUDA, DAN OLAHRAGA', 'SLB NEGERI 1 KULON PROGO', 'Jl. Srikandi, Pengasih, Kulon Progo, DIY', 'Dra. Hj. ...', '19670101...', ?, '19900101...', 'Kulon Progo, 18 Agustus 2026')
                    """,
              (
                  reg_user.strip(),
                  hash_password(reg_pass),
                  reg_name.strip(),
                  reg_name.strip(),
              ),
          )
          conn.commit()
          st.success("Akun berhasil dibuat! Silakan login di panel kiri.")
        except sqlite3.IntegrityError:
          st.error("Username sudah terdaftar! Gunakan username lain.")
        finally:
          conn.close()


# ==========================================
# 5. MODUL UTAMA APLIKASI (DASHBOARD GURU)
# ==========================================
def main_dashboard():
  user_row = get_user_profile(st.session_state.username)

  st.markdown(
      f"""
        <div class="main-header">
            <h1>🏫 LANTIP AI</h1>
            <p>Selamat Datang, <b>{user_row[2] or st.session_state.username}</b>! (Platform Lembar Administrasi Pembelajaran Inklusif Terintegrasi AI)</p>
        </div>
    """,
      unsafe_allow_html=True,
  )

  # --- SIDEBAR ---
  with st.sidebar:
    st.header(f"👤 Profil Guru: {st.session_state.username}")
    if st.button("🚪 Keluar (Logout)"):
      st.session_state.logged_in = False
      st.session_state.username = ""
      st.rerun()

    st.divider()
    st.header("⚙️ Pengaturan Tersimpan (Auto-Save)")

    api_key_db = st.text_input(
        "Gemini API Key",
        value=user_row[3] or "",
        type="password",
        help="API Key tersimpan otomatis di akun Anda",
    )
    nama_dinas_db = st.text_input("Nama Dinas", value=user_row[4] or "")
    nama_sekolah_db = st.text_input("Nama Sekolah", value=user_row[5] or "")
    alamat_sekolah_db = st.text_input("Alamat Sekolah", value=user_row[6] or "")

    st.caption("Penandatangan Dokumen:")
    nama_kepsek_db = st.text_input(
        "Nama Kepala Sekolah", value=user_row[7] or ""
    )
    nip_kepsek_db = st.text_input("NIP Kepala Sekolah", value=user_row[8] or "")
    nama_guru_db = st.text_input("Nama Guru Pengampu", value=user_row[9] or "")
    nip_guru_db = st.text_input("NIP Guru Pengampu", value=user_row[10] or "")
    kota_tgl_db = st.text_input(
        "Tempat & Tanggal Cetak", value=user_row[11] or ""
    )

    if st.button("💾 Simpan Perubahan Profil"):
      p_data = {
          "api_key": api_key_db.strip(),
          "nama_dinas": nama_dinas_db.strip(),
          "nama_sekolah": nama_sekolah_db.strip(),
          "alamat_sekolah": alamat_sekolah_db.strip(),
          "nama_kepsek": nama_kepsek_db.strip(),
          "nip_kepsek": nip_kepsek_db.strip(),
          "nama_guru": nama_guru_db.strip(),
          "nip_guru": nip_guru_db.strip(),
          "kota_tgl": kota_tgl_db.strip(),
      }
      update_user_profile(st.session_state.username, p_data)
      st.success("✅ Profil & API Key berhasil disimpan permanen!")
      st.rerun()

  # NAVIGASI TAB UTAMA
  nav_tab1, nav_tab2 = st.tabs(
      ["⚡ Workspace LANTIP AI", "📁 Riwayat Pekerjaan Saya"]
  )

  # ==========================================
  # TAB 1: WORKSPACE AI GENERATOR
  # ==========================================
  with nav_tab1:
    st.subheader("🎯 Langkah 1: Input Capaian Pembelajaran (CP)")

    col_a, col_b, col_c = st.columns(3)
    with col_a:
      mata_pelajaran = st.selectbox(
          "Mata Pelajaran",
          ["Matematika", "Bahasa Indonesia", "IPAS", "Seni Budaya"],
      )
    with col_b:
      fase_kelas = st.selectbox(
          "Fase / Kelas",
          ["Fase A / Kelas I", "Fase A / Kelas III", "Fase B / Kelas IV"],
      )
    with col_c:
      kekhususan = st.selectbox(
          "Kekhususan Siswa SLB",
          [
              "Hambatan Intelektual",
              "Hambatan Pendengaran",
              "Hambatan Penglihatan",
              "Autisme",
          ],
      )

    cp_input = st.text_area(
        "Tempelkan Teks Capaian Pembelajaran (CP) Kurikulum Merdeka:",
        "Mengurutkan dan membandingkan banyak-sedikit dengan benda konkret"
        " sampai dengan 10 serta memahami besar-kecil suatu benda",
        height=90,
    )

    if st.button(
        "🤖 Generasi Otomatis Administrasi Komplit (CP → TP → ATP → Prota →"
        " Prosem → Modul Ajar → PPI → Rubrik)",
        type="primary",
    ):
      api_key_use = api_key_db.strip()
      if not api_key_use:
        st.error(
            "⚠️ API Key belum diisi! Silakan masukkan Gemini API Key Anda pada"
            " menu Sidebar kiri lalu klik 'Simpan Perubahan Profil'."
        )
      else:
        with st.spinner(
            "LANTIP AI sedang merumuskan CP-TP-ATP, Prota-Prosem, Modul Ajar (8"
            " Dimensi Profil Lulusan), PPI & Rubrik Asesmen..."
        ):
          try:
            client = genai.Client(api_key=api_key_use)

            prompt = f"""
                        Kamu adalah pakar kurikulum pembelajaran inklusif SLB di Indonesia (LANTIP AI - Lembar Administrasi Pembelajaran).
                        Tolong analisis Capaian Pembelajaran (CP) berikut secara mendalam untuk siswa SLB.
                        Kekhususan: {kekhususan}, Fase/Kelas: {fase_kelas}, Mata Pelajaran: {mata_pelajaran}.

                        Teks CP: "{cp_input}"

                        INSTRUKSI KHUSUS LANTIP AI:
                        1. Integrasikan 8 Dimensi Profil Lulusan yang paling relevan.
                        2. Buatkan rekomendasi Program Pembelajaran Individual (PPI) meliputi Baseline, Akomodasi, serta Target Jangka Panjang & Jangka Pendek.
                        3. Buatkan Rubrik Asesmen Adaptif (Ceklis Observasi, Skala Penilaian, & Catatan Anekdot).

                        Berikan respon HANYA dalam format teks terstruktur persis seperti pola di bawah ini (gunakan pemisah tanda titik dua ':'):

                        ELEMEN: [Nama Elemen CP]
                        RUANG LINGKUP: [Rincian Ruang Lingkup Materi]
                        MATERI POKOK: [Daftar Materi Pokok]
                        TP: [Jabarkan Poin 1, 2, 3... Tujuan Pembelajaran Adaptif lengkap]
                        ATP1: [ATP Tahap 1 Pengenalan/Konkret]
                        ATP2: [ATP Tahap 2 Pemahaman/Koneksi]
                        ATP3: [ATP Tahap 3 Penerapan/Respon Mandiri]
                        KEGIATAN: [Kegiatan Konkret Adaptif SLB]
                        ASESMEN: [Bentuk Asesmen Adaptif / Observasi]
                        ALOKASI: [Total JP, misal: 24 JP]
                        PROTA_SEM1: [Daftar TP & Alokasi JP untuk Semester 1]
                        PROTA_SEM2: [Daftar TP & Alokasi JP untuk Semester 2]
                        PROSEM: [Distribusi Mingguan per Bulan untuk Sem 1 & 2]
                        DIMENSI_LULUSAN: [List Dimensi dari 8 Dimensi Profil Lulusan]
                        MODUL_AJAR: [Rancangan Modul Ajar Adaptif: Identitas, Target 8 Dimensi Lulusan, Media Konkret, Langkah Pembelajaran Deferensiasi, & Refleksi]
                        PPI_TEXT: [Dokumen PPI: Baseline Kemampuan Awal, Akomodasi Pembelajaran, Target Jangka Panjang, Target Jangka Pendek]
                        RUBRIK_TEXT: [Instrumen Asesmen: Ceklis Observasi, Skala Penilaian 1-3, dan Format Catatan Anekdot]
                        """

            response = client.models.generate_content(
                model="gemini-3.6-flash", contents=prompt
            )
            res_text = response.text
            parsed = {}
            for line in res_text.split("\n"):
              if ":" in line:
                key, val = line.split(":", 1)
                parsed[key.strip()] = val.strip()

            for key_map, state_key in [
                ("ELEMEN", "elemen"),
                ("RUANG LINGKUP", "ruang_lingkup"),
                ("MATERI POKOK", "materi_pokok"),
                ("TP", "tp_text"),
                ("ATP1", "atp1"),
                ("ATP2", "atp2"),
                ("ATP3", "atp3"),
                ("KEGIATAN", "kegiatan_konkret"),
                ("ASESMEN", "bentuk_asesmen"),
                ("ALOKASI", "alokasi_waktu"),
                ("PROTA_SEM1", "prota_sem1"),
                ("PROTA_SEM2", "prota_sem2"),
                ("PROSEM", "prosem_detail"),
                ("DIMENSI_LULUSAN", "dimensi_lulusan"),
                ("MODUL_AJAR", "modul_ajar"),
                ("PPI_TEXT", "ppi_text"),
                ("RUBRIK_TEXT", "rubrik_text"),
            ]:
              if key_map in parsed:
                st.session_state.form_data[state_key] = parsed[key_map]

            title_project = (
                f"Paket Administrasi {mata_pelajaran} {fase_kelas}"
                f" ({kekhususan})"
            )
            save_history(
                st.session_state.username,
                title_project,
                mata_pelajaran,
                fase_kelas,
                kekhususan,
                cp_input,
                st.session_state.form_data,
            )

            st.success(
                "✅ Paket Administrasi Komplit Berhasil Dibuat oleh LANTIP AI &"
                " Tersimpan Otomatis!"
            )
            st.rerun()
          except Exception as e:
            st.error(f"Pesan Error AI: {e}")

    st.divider()

    # REVIEW DOKUMEN HASIL ANALISIS
    st.subheader("📋 Langkah 2: Review & Unduh Paket Administrasi")

    t1, t2, t3, t4, t5, t6 = st.tabs([
        "📊 1. Analisis CP-TP-ATP",
        "📅 2. Prota",
        "🗓️ 3. Prosem",
        "📝 4. Modul Ajar (8 Dimensi)",
        "🩺 5. Dokumen PPI",
        "📋 6. Rubrik Asesmen & Ceklis",
    ])

    with t1:
      elemen = st.text_input("Elemen CP", st.session_state.form_data["elemen"])
      c1, c2 = st.columns(2)
      with c1:
        ruang_lingkup = st.text_area(
            "Ruang Lingkup Materi",
            st.session_state.form_data["ruang_lingkup"],
            height=90,
        )
      with c2:
        materi_pokok = st.text_area(
            "Materi Pokok",
            st.session_state.form_data["materi_pokok"],
            height=90,
        )

      tp_text = st.text_area(
          "Penjabaran Tujuan Pembelajaran (TP)",
          st.session_state.form_data["tp_text"],
          height=120,
      )

      st.write("**Alur Tujuan Pembelajaran (ATP) Berjenjang:**")
      ca1, ca2, ca3 = st.columns(3)
      with ca1:
        atp1 = st.text_area(
            "Tahap 1 (Konkret)",
            st.session_state.form_data["atp1"],
            height=100,
        )
      with ca2:
        atp2 = st.text_area(
            "Tahap 2 (Pemahaman)",
            st.session_state.form_data["atp2"],
            height=100,
        )
      with ca3:
        atp3 = st.text_area(
            "Tahap 3 (Penerapan)",
            st.session_state.form_data["atp3"],
            height=100,
        )

      cx, cy, cz = st.columns(3)
      with cx:
        kegiatan_konkret = st.text_area(
            "Kegiatan Konkret Adaptif",
            st.session_state.form_data["kegiatan_konkret"],
            height=90,
        )
      with cy:
        bentuk_asesmen = st.text_area(
            "Bentuk Asesmen Inklusif",
            st.session_state.form_data["bentuk_asesmen"],
            height=90,
        )
      with cz:
        alokasi_waktu = st.text_input(
            "Alokasi Waktu Total", st.session_state.form_data["alokasi_waktu"]
        )

    with t2:
      cp1, cp2 = st.columns(2)
      with cp1:
        prota_sem1 = st.text_area(
            "Program Tahunan Semester 1 (TP & JP)",
            st.session_state.form_data["prota_sem1"],
            height=180,
        )
      with cp2:
        prota_sem2 = st.text_area(
            "Program Tahunan Semester 2 (TP & JP)",
            st.session_state.form_data["prota_sem2"],
            height=180,
        )

    with t3:
      prosem_detail = st.text_area(
          "Distribusi Program Semester (Prosem)",
          st.session_state.form_data["prosem_detail"],
          height=200,
      )

    with t4:
      dimensi_lulusan = st.text_area(
          "Target 8 Dimensi Profil Lulusan yang Dikembangkan",
          st.session_state.form_data["dimensi_lulusan"],
          height=90,
      )
      modul_ajar = st.text_area(
          "Draft Modul Ajar Adaptif Rinci",
          st.session_state.form_data["modul_ajar"],
          height=200,
      )

    with t5:
      ppi_text = st.text_area(
          "Rancangan Program Pembelajaran Individual (PPI)",
          st.session_state.form_data["ppi_text"],
          height=260,
      )

    with t6:
      rubrik_text = st.text_area(
          "Instrumen Rubrik Asesmen & Ceklis Observasi Adaptif",
          st.session_state.form_data["rubrik_text"],
          height=260,
      )

    # --- FUNGSIONALITAS CETAK WORD DOCX ---
    def generate_docx():
      doc = docx.Document()
      sec = doc.sections[0]
      sec.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
      sec.page_width = Inches(11.69)
      sec.page_height = Inches(8.27)
      sec.top_margin = Inches(0.5)
      sec.bottom_margin = Inches(0.5)
      sec.left_margin = Inches(0.6)
      sec.right_margin = Inches(0.6)

      def set_cell_background(cell, fill_hex):
        tcPr = cell._element.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

      def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
        tcPr = cell._element.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [
            ('top', top),
            ('bottom', bottom),
            ('left', left),
            ('right', right),
        ]:
          node = OxmlElement(f'w:{m}')
          node.set(qn('w:w'), str(val))
          node.set(qn('w:type'), 'dxa')
          tcMar.append(node)
        tcPr.append(tcMar)

      # Kop Surat
      kop_tbl = doc.add_table(rows=1, cols=2)
      kop_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
      c_logo, c_info = kop_tbl.rows[0].cells
      c_logo.width = Inches(1.2)
      c_info.width = Inches(9.29)

      c_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
      c_logo.paragraphs[0].add_run("[ LOGO ]").font.size = Pt(8)

      p_info = c_info.paragraphs[0]
      p_info.add_run(f'{nama_dinas_db.upper()}\n').font.size = Pt(9)
      r_sch = p_info.add_run(f'{nama_sekolah_db.upper()}\n')
      r_sch.font.size = Pt(13)
      r_sch.font.bold = True
      p_info.add_run(f'Alamat: {alamat_sekolah_db}').font.size = Pt(8)

      p_hr = doc.add_paragraph()
      p_hr.paragraph_format.space_before = Pt(4)
      p_hr.paragraph_format.space_after = Pt(8)
      p_hr._element.get_or_add_pPr().append(
          parse_xml(
              f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12"'
              ' w:space="1" w:color="1B365D"/></w:pBdr>'
          )
      )

      # Judul
      p_title = doc.add_paragraph()
      p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
      r_t = p_title.add_run(
          "PAKET PERENCANAAN PEMBELAJARAN INKLUSIF & ADMINISTRASI GURU SLB"
      )
      r_t.font.bold = True
      r_t.font.size = Pt(12)
      r_t.font.color.rgb = RGBColor(27, 54, 93)

      p_sub = doc.add_paragraph()
      p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
      p_sub.paragraph_format.space_after = Pt(10)
      p_sub.add_run(
          f"Mata Pelajaran: {mata_pelajaran} | {fase_kelas} | Kekhususan:"
          f" {kekhususan}"
      ).font.size = Pt(9)

      # Tabel Analisis
      headers = [
          "No",
          "Elemen",
          "Capaian Pembelajaran (CP)",
          "Ruang Lingkup Materi",
          "Materi Pokok",
          "Tujuan Pembelajaran (TP)",
          "Alur Tujuan Pembelajaran (ATP)",
          "Kegiatan Konkret",
          "Bentuk Asesmen",
          "Alokasi Waktu",
      ]
      col_widths = [0.4, 0.9, 1.4, 1.0, 0.9, 1.3, 1.6, 1.4, 0.9, 0.69]

      table = doc.add_table(rows=1, cols=10)
      table.alignment = WD_TABLE_ALIGNMENT.CENTER

      for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(col_widths[idx])
        set_cell_background(cell, "1B365D")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)

      row_cells = table.add_row().cells
      atp_combined = (
          f"Tahap 1:\n{atp1}\n\nTahap 2:\n{atp2}\n\nTahap 3:\n{atp3}"
      )
      data = [
          "1",
          elemen,
          cp_input,
          ruang_lingkup,
          materi_pokok,
          tp_text,
          atp_combined,
          kegiatan_konkret,
          bentuk_asesmen,
          alokasi_waktu,
      ]

      for idx, val in enumerate(data):
        row_cells[idx].width = Inches(col_widths[idx])
        set_cell_background(row_cells[idx], "FFFFFF")
        set_cell_margins(row_cells[idx])
        p = row_cells[idx].paragraphs[0]
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if idx in [0, 9]
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        run = p.add_run(val)
        run.font.size = Pt(7.5)

      doc.add_paragraph().paragraph_format.space_before = Pt(12)

      # Lampiran
      p_pro = doc.add_paragraph()
      p_pro.add_run(
          "LAMPIRAN PROTA, PROSEM, MODUL AJAR, PPI & RUBRIK ASESMEN:\n"
      ).font.bold = True
      p_pro.add_run(
          f"• Semester 1:\n{prota_sem1}\n\n• Semester 2:\n{prota_sem2}\n\n• Target"
          f" 8 Dimensi Profil Lulusan:\n{dimensi_lulusan}\n\n• Modul Ajar"
          f" Adaptif:\n{modul_ajar}\n\n• Program Pembelajaran Individual"
          f" (PPI):\n{ppi_text}\n\n• Rubrik Asesmen & Ceklis"
          f" Observasi:\n{rubrik_text}"
      ).font.size = Pt(8)

      doc.add_paragraph().paragraph_format.space_before = Pt(12)

      # Tanda Tangan
      sig_tbl = doc.add_table(rows=1, cols=2)
      sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
      cL, cR = sig_tbl.rows[0].cells
      cL.width = Inches(5.2)
      cR.width = Inches(5.2)

      cL.paragraphs[0].add_run(
          f"Mengetahui,\nKepala"
          f" {nama_sekolah_db}\n\n\n\n{nama_kepsek_db}\nNIP. {nip_kepsek_db}"
      ).font.size = Pt(8.5)
      cR.paragraphs[0].add_run(
          f"{kota_tgl_db}\nGuru Kelas /"
          f" Pengampu,\n\n\n\n{nama_guru_db}\nNIP. {nip_guru_db}"
      ).font.size = Pt(8.5)

      bio = io.BytesIO()
      doc.save(bio)
      return bio.getvalue()

    st.markdown("---")
    col_d1, col_d2 = st.columns([2, 1])
    with col_d1:
      st.write("🚀 **Dokumen Siap Dicetak:**")
      st.caption(
          "Unduh paket komplit dokumen administrasi (CP-TP-ATP, Prota, Prosem,"
          " Modul Ajar 8 Dimensi Lulusan, PPI, dan Rubrik Asesmen) dalam format"
          " Word (.docx)."
      )
    with col_d2:
      docx_bytes = generate_docx()
      st.download_button(
          label="📥 Download File Word Komplit (.docx)",
          data=docx_bytes,
          file_name=(
              f"Administrasi_LANTIP_{mata_pelajaran}_{st.session_state.username}.docx"
          ),
          mime=(
              "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          ),
          type="primary",
      )

  # ==========================================
  # TAB 2: RIWAYAT PEKERJAAN
  # ==========================================
  with nav_tab2:
    st.subheader("📁 Database Riwayat Pekerjaan Saya")
    history_items = get_user_history(st.session_state.username)

    if not history_items:
      st.info(
          "Belum ada riwayat dokumen yang tersimpan. Gunakan menu 'Workspace"
          " LANTIP AI' untuk membuat analisis CP baru!"
      )
    else:
      for item in history_items:
        (
            item_id,
            created_at,
            title,
            mapel,
            fase_kelas,
            kekhususan,
            cp_in,
            data_json_str,
        ) = item

        with st.expander(
            f"📌 [{created_at}] {title} - Mapel: {mapel} ({kekhususan})"
        ):
          st.write(f"**Teks CP Asli:** {cp_in}")

          c_h1, c_h2 = st.columns([1, 1])
          with c_h1:
            if st.button(f"📥 Buka & Load ke Workspace", key=f"load_{item_id}"):
              st.session_state.form_data = json.loads(data_json_str)
              st.success(
                  "Dokumen berhasil dimuat ke Workspace! Silakan cek di tab"
                  " 'Workspace LANTIP AI'."
              )
              st.rerun()

          with c_h2:
            if st.button(f"🗑️ Hapus dari Riwayat", key=f"del_{item_id}"):
              delete_history_item(item_id)
              st.success("Dokumen telah dihapus dari riwayat.")
              st.rerun()


# ==========================================
# 6. ROUTING UTAMA
# ==========================================
if st.session_state.logged_in:
  main_dashboard()
else:
  login_page()
